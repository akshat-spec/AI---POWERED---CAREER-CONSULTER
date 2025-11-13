from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_project_report():
    doc = Document()

    # Title Page
    title = doc.add_heading('Intelligent Career Guidance System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Project Report', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Prepared by: Data Scientist')
    doc.add_paragraph('Date: 2024')

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc = doc.add_paragraph()
    toc.add_run('1. Introduction\n')
    toc.add_run('2. Project Overview\n')
    toc.add_run('3. Technical Architecture\n')
    toc.add_run('4. Machine Learning Model\n')
    toc.add_run('5. Key Features and Functionality\n')
    toc.add_run('6. How Things Work\n')
    toc.add_run('7. Technologies Used\n')
    toc.add_run('8. Deployment and Usage\n')
    toc.add_run('9. Future Enhancements\n')
    toc.add_run('10. Conclusion\n')

    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', 1)
    intro = doc.add_paragraph()
    intro.add_run('As a data scientist specializing in machine learning applications for web-based solutions, I have developed a comprehensive Intelligent Career Guidance System. This project leverages machine learning to provide personalized career recommendations based on user skill assessments. The system is built as a full-stack web application using Flask for the backend and integrates a pre-trained machine learning model for career prediction.')

    # 2. Project Overview
    doc.add_heading('2. Project Overview', 1)

    doc.add_heading('2.1 Objectives', 2)
    objectives = doc.add_paragraph()
    objectives.add_run('• Create a user-friendly web platform for career guidance\n')
    objectives.add_run('• Implement machine learning-based career prediction\n')
    objectives.add_run('• Provide a seamless user experience with authentication\n')
    objectives.add_run('• Offer additional resources like courses, blogs, and contact information\n')

    doc.add_heading('2.2 Key Features', 2)
    features = doc.add_paragraph()
    features.add_run('• User registration and authentication\n')
    features.add_run('• Skill assessment form with 17 different skill categories\n')
    features.add_run('• Machine learning-powered career prediction\n')
    features.add_run('• Multiple career recommendations with probability scores\n')
    features.add_run('• Responsive web interface\n')
    features.add_run('• Navigation to additional resources\n')

    # 3. Technical Architecture
    doc.add_heading('3. Technical Architecture', 1)

    doc.add_heading('3.1 Backend (Flask Application)', 2)
    backend = doc.add_paragraph()
    backend.add_run('The backend is built using Python Flask framework with the following components:')

    doc.add_heading('3.1.1 Routes and Endpoints', 3)
    routes = doc.add_paragraph()
    routes.add_run('• / (GET): Home page (redirects to signup if not logged in)\n')
    routes.add_run('• /signup (GET/POST): User registration\n')
    routes.add_run('• /login (GET/POST): User authentication\n')
    routes.add_run('• /logout (GET): User logout\n')
    routes.add_run('• /predict (POST): Career prediction processing\n')
    routes.add_run('• /about (GET): About page\n')
    routes.add_run('• /contact (GET): Contact page\n')
    routes.add_run('• /blog (GET): Knowledge network page\n')
    routes.add_run('• /courses (GET): Courses page\n')

    doc.add_heading('3.1.2 Authentication System', 3)
    auth = doc.add_paragraph()
    auth.add_run('• Session-based authentication using Flask sessions\n')
    auth.add_run('• In-memory user storage (for demonstration; should be replaced with database in production)\n')
    auth.add_run('• Automatic login after successful registration\n')
    auth.add_run('• Protected routes requiring authentication\n')

    doc.add_heading('3.1.3 Machine Learning Integration', 3)
    ml = doc.add_paragraph()
    ml.add_run('• Pre-trained scikit-learn model loaded from careerlast.pkl\n')
    ml.add_run('• Model expects 17 input features corresponding to skill ratings\n')
    ml.add_run('• Uses predict() for primary career recommendation\n')
    ml.add_run('• Uses predict_proba() for probability scores of all careers\n')
    ml.add_run('• Returns top matching careers based on probability thresholds\n')

    doc.add_heading('3.2 Frontend (HTML Templates)', 2)
    frontend = doc.add_paragraph()
    frontend.add_run('The frontend uses Jinja2 templating with Bootstrap for styling:')

    doc.add_heading('3.2.1 Templates', 3)
    templates = doc.add_paragraph()
    templates.add_run('• signup.html: User registration form\n')
    templates.add_run('• login.html: User login form\n')
    templates.add_run('• hometest.html: Skill assessment form\n')
    templates.add_run('• testafter.html: Career prediction results\n')
    templates.add_run('• about.html: About the system\n')
    templates.add_run('• contact.html: Contact information\n')
    templates.add_run('• blog.html: Knowledge network\n')
    templates.add_run('• courses.html: Available courses\n')

    doc.add_heading('3.2.2 User Interface Features', 3)
    ui = doc.add_paragraph()
    ui.add_run('• Responsive design using Bootstrap CSS\n')
    ui.add_run('• Form validation for required fields\n')
    ui.add_run('• Dropdown menus for skill rating (1-9 scale)\n')
    ui.add_run('• Dynamic career recommendation display\n')
    ui.add_run('• Navigation bar with dropdown menus\n')

    doc.add_heading('3.3 Data Flow', 2)
    dataflow = doc.add_paragraph()
    dataflow.add_run('1. User Registration/Login: User creates account or logs in\n')
    dataflow.add_run('2. Skill Assessment: User rates themselves on 17 skills (1-9 scale)\n')
    dataflow.add_run('3. Data Processing: Form data converted to numpy array\n')
    dataflow.add_run('4. Model Prediction: ML model predicts primary career and probabilities\n')
    dataflow.add_run('5. Result Display: Top matching careers shown with links to detailed pages\n')
    dataflow.add_run('6. Navigation: User can explore additional resources\n')

    # 4. Machine Learning Model
    doc.add_heading('4. Machine Learning Model', 1)

    doc.add_heading('4.1 Model Details', 2)
    model_details = doc.add_paragraph()
    model_details.add_run('• Type: Classification model (likely Random Forest or similar ensemble method)\n')
    model_details.add_run('• Input Features: 17 skill ratings (Database Fundamentals, Computer Architecture, etc.)\n')
    model_details.add_run('• Output: 17 career categories\n')
    model_details.add_run('• Training Data: Based on dataset9000.csv and dataset9000.data\n')
    model_details.add_run('• Model File: careerlast.pkl (serialized scikit-learn model)\n')

    doc.add_heading('4.2 Prediction Logic', 2)
    prediction_logic = doc.add_paragraph()
    prediction_logic.add_run('```python\n')
    prediction_logic.add_run('# Load model\n')
    prediction_logic.add_run('loaded_model = pickle.load(open("careerlast.pkl", \'rb\'))\n\n')
    prediction_logic.add_run('# Prepare input\n')
    prediction_logic.add_run('data = np.array(arr).reshape(1, -1)\n\n')
    prediction_logic.add_run('# Primary prediction\n')
    prediction_logic.add_run('predictions = loaded_model.predict(data)\n\n')
    prediction_logic.add_run('# Probability scores\n')
    prediction_logic.add_run('pred = loaded_model.predict_proba(data)\n')
    prediction_logic.add_run('pred = pred > 0.05  # Threshold for secondary recommendations\n\n')
    prediction_logic.add_run('# Filter secondary careers\n')
    prediction_logic.add_run('for i in range(17):\n')
    prediction_logic.add_run('    if pred[0, i]:\n')
    prediction_logic.add_run('        final_res.append(i)\n')
    prediction_logic.add_run('```\n')

    doc.add_heading('4.3 Career Mapping', 2)
    career_mapping = doc.add_paragraph()
    career_mapping.add_run('The system maps numeric predictions to career names:\n')
    career_mapping.add_run('• 0: AI ML Specialist\n')
    career_mapping.add_run('• 1: API Integration Specialist\n')
    career_mapping.add_run('• 2: Application Support Engineer\n')
    career_mapping.add_run('• ... (17 total careers)\n')

    # 5. Key Features and Functionality
    doc.add_heading('5. Key Features and Functionality', 1)

    doc.add_heading('5.1 User Authentication', 2)
    user_auth = doc.add_paragraph()
    user_auth.add_run('• Secure session management\n')
    user_auth.add_run('• Password-based authentication\n')
    user_auth.add_run('• Automatic redirection for unauthenticated users\n')
    user_auth.add_run('• Logout functionality\n')

    doc.add_heading('5.2 Career Prediction Algorithm', 2)
    algorithm = doc.add_paragraph()
    algorithm.add_run('1. User inputs skill ratings (1-9 scale for each of 17 skills)\n')
    algorithm.add_run('2. Data normalized and fed to ML model\n')
    algorithm.add_run('3. Primary career recommended\n')
    algorithm.add_run('4. Secondary careers suggested based on probability > 5%\n')
    algorithm.add_run('5. Results displayed with visual styling\n')

    doc.add_heading('5.3 Web Interface', 2)
    web_interface = doc.add_paragraph()
    web_interface.add_run('• Clean, professional design\n')
    web_interface.add_run('• Mobile-responsive layout\n')
    web_interface.add_run('• Intuitive navigation\n')
    web_interface.add_run('• Form validation\n')
    web_interface.add_run('• Error handling for invalid inputs\n')

    doc.add_heading('5.4 Additional Resources', 2)
    resources = doc.add_paragraph()
    resources.add_run('• About page explaining the system\n')
    resources.add_run('• Contact page for inquiries\n')
    resources.add_run('• Blog/Knowledge Network for career information\n')
    resources.add_run('• Courses page for learning resources\n')

    # 6. How Things Work
    doc.add_heading('6. How Things Work', 1)

    doc.add_heading('6.1 User Journey', 2)
    user_journey = doc.add_paragraph()
    user_journey.add_run('1. Access: User visits http://127.0.0.1:5000\n')
    user_journey.add_run('2. Registration: New users sign up with username/password\n')
    user_journey.add_run('3. Login: Existing users authenticate\n')
    user_journey.add_run('4. Assessment: Rate skills on 17 categories\n')
    user_journey.add_run('5. Prediction: ML model analyzes input and predicts careers\n')
    user_journey.add_run('6. Results: View primary and secondary career recommendations\n')
    user_journey.add_run('7. Exploration: Navigate to detailed career pages or other resources\n')

    doc.add_heading('6.2 Data Processing Pipeline', 2)
    pipeline = doc.add_paragraph()
    pipeline.add_run('User Input → Form Validation → Data Conversion → ML Model → Prediction → Template Rendering → User Display\n')

    doc.add_heading('6.3 Session Management', 2)
    session = doc.add_paragraph()
    session.add_run('• Flask sessions store user authentication state\n')
    session.add_run('• Protected routes check session before allowing access\n')
    session.add_run('• Logout clears session data\n')

    doc.add_heading('6.4 Static File Serving', 2)
    static = doc.add_paragraph()
    static.add_run('• CSS, JS, and images served from /static/ directory\n')
    static.add_run('• Templates use url_for() for dynamic URL generation\n')
    static.add_run('• Bootstrap and custom styles for responsive design\n')

    # 7. Technologies Used
    doc.add_heading('7. Technologies Used', 1)

    doc.add_heading('7.1 Backend', 2)
    backend_tech = doc.add_paragraph()
    backend_tech.add_run('• Python 3.x: Core programming language\n')
    backend_tech.add_run('• Flask 2.3.3: Web framework\n')
    backend_tech.add_run('• NumPy 1.24.3: Numerical computing\n')
    backend_tech.add_run('• Scikit-learn 1.3.0: Machine learning library\n')

    doc.add_heading('7.2 Frontend', 2)
    frontend_tech = doc.add_paragraph()
    frontend_tech.add_run('• HTML5: Markup language\n')
    frontend_tech.add_run('• CSS3: Styling (Bootstrap + custom)\n')
    frontend_tech.add_run('• JavaScript: Client-side interactivity\n')
    frontend_tech.add_run('• Jinja2: Template engine\n')

    doc.add_heading('7.3 Development Tools', 2)
    dev_tools = doc.add_paragraph()
    dev_tools.add_run('• VS Code: IDE\n')
    dev_tools.add_run('• Git: Version control\n')
    dev_tools.add_run('• Pip: Package management\n')

    # 8. Deployment and Usage
    doc.add_heading('8. Deployment and Usage', 1)

    doc.add_heading('8.1 Local Development', 2)
    local_dev = doc.add_paragraph()
    local_dev.add_run('```bash\n')
    local_dev.add_run('pip install -r requirements.txt\n')
    local_dev.add_run('python testapp.py\n')
    local_dev.add_run('```\n')
    local_dev.add_run('Access at: http://127.0.0.1:5000\n')

    doc.add_heading('8.2 Production Considerations', 2)
    production = doc.add_paragraph()
    production.add_run('• Replace in-memory user storage with database (SQLite/PostgreSQL)\n')
    production.add_run('• Use secure session keys\n')
    production.add_run('• Implement password hashing\n')
    production.add_run('• Add input sanitization\n')
    production.add_run('• Use WSGI server (Gunicorn) instead of development server\n')

    # 9. Future Enhancements
    doc.add_heading('9. Future Enhancements', 1)

    doc.add_heading('9.1 Technical Improvements', 2)
    tech_improvements = doc.add_paragraph()
    tech_improvements.add_run('• Database integration for user management\n')
    tech_improvements.add_run('• Advanced ML models (deep learning)\n')
    tech_improvements.add_run('• Real-time prediction feedback\n')
    tech_improvements.add_run('• User profile and history tracking\n')

    doc.add_heading('9.2 Feature Additions', 2)
    feature_additions = doc.add_paragraph()
    feature_additions.add_run('• Career roadmap generation\n')
    feature_additions.add_run('• Skill gap analysis\n')
    feature_additions.add_run('• Job market integration\n')
    feature_additions.add_run('• Social features (sharing results)\n')
    feature_additions.add_run('• Admin panel for content management\n')

    doc.add_heading('9.3 Data Science Enhancements', 2)
    ds_enhancements = doc.add_paragraph()
    ds_enhancements.add_run('• Larger training datasets\n')
    ds_enhancements.add_run('• Feature engineering for better predictions\n')
    ds_enhancements.add_run('• Model explainability (SHAP values)\n')
    ds_enhancements.add_run('• A/B testing for recommendation algorithms\n')

    # 10. Conclusion
    doc.add_heading('10. Conclusion', 1)
    conclusion = doc.add_paragraph()
    conclusion.add_run('This Intelligent Career Guidance System demonstrates the practical application of machine learning in web development. By combining a trained classification model with a user-friendly Flask application, the system provides valuable career insights based on skill assessments. The project showcases end-to-end data science workflow from model training to web deployment, offering a scalable solution for career guidance that can be extended with additional features and improved models.\n\n')
    conclusion.add_run('The system successfully integrates authentication, form handling, machine learning prediction, and responsive web design into a cohesive application that serves both new and experienced users in their career exploration journey.')

    # Save the document
    doc.save('Intelligent_Career_Guidance_System_Project_Report.docx')
    print("Project report saved as 'Intelligent_Career_Guidance_System_Project_Report.docx'")

if __name__ == "__main__":
    create_project_report()
